import logging

import pytest

pd = pytest.importorskip("pandas")

import raffalib.pandas  # noqa: E402,F401  registers the `raffa` accessor


def test_add_prefix_if_not_exists():
    df = pd.DataFrame({"a": [1], "pre_b": [2]})
    out = df.raffa.add_prefix_if_not_exists("pre_")
    assert list(out.columns) == ["pre_a", "pre_b"]


def test_get_duplicates():
    df = pd.DataFrame({"a": [1, 1, 2], "b": [1, 1, 3]})
    dups = df.raffa.get_duplicates()
    assert len(dups) == 1
    assert dups.iloc[0]["a"] == 1


def test_sort_columns_uses_natural_order():
    df = pd.DataFrame({"c10": [1], "c2": [1], "c1": [1]})
    out = df.raffa.sort_columns()
    assert list(out.columns) == ["c1", "c2", "c10"]


def test_series_freq():
    s = pd.Series(["a", "a", "b"], name="letter")
    f = s.raffa.freq()
    assert f.loc["a", "count"] == 2
    assert f.loc["b", "count"] == 1
    assert f.loc["Total", "count"] == 3
    assert f.index.name == "letter"


def test_dataframe_freq_delegates_to_series():
    df = pd.DataFrame({"letter": ["a", "a", "b"]})
    f = df.raffa.freq("letter")
    assert f.loc["Total", "count"] == 3


def test_endlog_logs_removed_rows(caplog):
    df = pd.DataFrame({"a": [1, 2, 3, 4]})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        df.raffa.startlog()
        df = df[df["a"] > 1]
        df.raffa.endlog()
    assert any(
        "Removed 1/4 (25.00%) rows. New shape: (3, 1)." in r.message
        for r in caplog.records
    )


def test_endlog_logs_value_changes(caplog):
    df = pd.DataFrame({"a": [1, 2, 3, 4]})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        df.raffa.startlog(clone=True)
        df = df.assign(a=[1, 99, 3, 99])
        df.raffa.endlog()
    assert any("Changed 2/4 (50.00%) values." in r.message for r in caplog.records)


def test_endlog_clone_false_message(caplog):
    df = pd.DataFrame({"a": [1, 2, 3]})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        df.raffa.startlog(clone=False)
        df.raffa.endlog()
    assert any("No value-level comparison" in r.message for r in caplog.records)


def test_series_endlog_removed_values(caplog):
    s = pd.Series([1, 2, 3, 4, 5])
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        s.raffa.startlog()
        s = s[s > 2]
        s.raffa.endlog()
    assert any(
        "Removed 2/5 (40.00%) values. New shape: (3,)." in r.message
        for r in caplog.records
    )


def test_join_inner_drops_source_columns(caplog):
    left = pd.DataFrame({"k": [1, 2, 3], "lv": [10, 20, 30]})
    right = pd.DataFrame({"k": [2, 3, 4], "rv": [200, 300, 400]})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        out = left.raffa.join(right, on="k", how="inner")
    assert out.shape[0] == 2
    assert set(out.columns) == {"k", "lv", "rv"}


def test_join_keep_row_index_appends_source_columns():
    left = pd.DataFrame({"k": [1, 2, 3], "lv": [10, 20, 30]})
    right = pd.DataFrame({"k": [2, 3, 4], "rv": [200, 300, 400]})
    out = left.raffa.join(right, on="k", how="inner", keep_row_index=True)
    assert list(out.columns)[-2:] == ["source_left", "source_right"]


def test_join_left_logs_row_provenance(caplog):
    left = pd.DataFrame({"k": [1, 2, 3, 4], "lv": [10, 20, 30, 40]})
    right = pd.DataFrame({"k": [2, 3], "rv": [200, 300]})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        out = left.raffa.join(right, on="k", how="left")
    assert out.shape[0] == 4
    log = "\n".join(r.message for r in caplog.records)
    assert "From left only: 2/4 (50.00%)" in log
    assert "From both: 2/4 (50.00%)" in log
    assert "Dropped rows from left: 0/4 (0.00%)" in log


def test_join_semi_is_detected_as_filtering(caplog):
    left = pd.DataFrame({"k": [1, 2, 3, 4, 5]})
    right = pd.DataFrame({"k": [2, 3, 4]})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        out = left.raffa.join(right, on="k", how="semi")
    assert out["k"].tolist() == [2, 3, 4]
    assert list(out.columns) == ["k"]
    assert any("Detected filtering join" in r.message for r in caplog.records)


def test_join_anti_returns_unmatched_left_rows():
    left = pd.DataFrame({"k": [1, 2, 3, 4, 5]})
    right = pd.DataFrame({"k": [2, 3, 4]})
    out = left.raffa.join(right, on="k", how="anti")
    assert out["k"].tolist() == [1, 5]
    assert list(out.columns) == ["k"]


def test_endlog_add_rows_to_empty_frame_is_guarded(caplog):
    df = pd.DataFrame({"a": []})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        df.raffa.startlog()
        df = df.reindex(range(3))
        df.raffa.endlog(timeit=False)
    assert any(
        "Added 3/0 (N/A) rows. New shape: (3, 1)." in r.message for r in caplog.records
    )


def test_join_empty_output_does_not_divide_by_zero(caplog):
    left = pd.DataFrame({"k": [1, 2]})
    right = pd.DataFrame({"k": [3, 4]})
    with caplog.at_level(logging.INFO, logger="raffalib.pandas"):
        out = left.raffa.join(right, on="k", how="inner")
    assert out.shape[0] == 0
    assert any("From both: 0/0 (N/A)" in r.message for r in caplog.records)


def test_to_docx_writes_file(tmp_path):
    pytest.importorskip("docx")
    df = pd.DataFrame({"a": [1, 2], "b": [3, 4]})
    out = tmp_path / "df.docx"
    df.raffa.to_docx(out)
    assert out.exists()


def test_to_docx_applies_heading_and_table_options(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    df = pd.DataFrame({"a": [1, 2], "b": [3, 4]})
    out = tmp_path / "styled.docx"
    df.raffa.to_docx(
        out, include_index=False, heading_text="Table 1", table_style="Light Grid"
    )
    reopened = Document(str(out))
    assert "Table 1" in [p.text for p in reopened.paragraphs]
    assert reopened.tables[0].style.name == "Light Grid"
