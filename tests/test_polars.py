import logging

import pytest

pl = pytest.importorskip("polars")
pytest.importorskip("polars_config_meta")

import raffalib.polars  # noqa: E402,F401  registers the `raffa` namespace
from raffalib.polars import PercOptions  # noqa: E402


def test_series_freq_has_total_row():
    s = pl.Series("x", ["a", "a", "b"])
    f = s.raffa.freq()
    assert f.columns == ["value", "count", "perc"]
    assert f["value"].to_list()[-1] == "Total"
    total_count = f.filter(pl.col("value") == "Total")["count"][0]
    assert total_count == 3


def test_dataframe_crosstab_counts():
    df = pl.DataFrame({"a": ["x", "x", "y"], "b": ["m", "n", "m"]})
    ct = df.raffa.crosstab("a", "b")
    assert ct.shape[0] == 2  # two distinct values of 'a'
    assert "a" in ct.columns


def test_dataframe_crosstab_row_percentages():
    df = pl.DataFrame({"a": ["x", "x", "y"], "b": ["m", "n", "m"]})
    ct = df.raffa.crosstab("a", "b", perc=PercOptions.ROWS)
    # row 'x' has one 'm' and one 'n' -> 50% each
    row_x = ct.filter(pl.col("a") == "x")
    assert row_x["m"][0] == pytest.approx(50.0)
    assert row_x["n"][0] == pytest.approx(50.0)


def test_endlog_removed_rows(caplog):
    df = pl.DataFrame({"a": [1, 2, 3, 4]})
    with caplog.at_level(logging.INFO, logger="raffalib.polars"):
        df.raffa.startlog()
        df = df.filter(pl.col("a") > 1)
        df.raffa.endlog(timeit=False)
    assert any("Removed 1/4 (25.00%) rows." in r.message for r in caplog.records)


def test_series_endlog_matches_pandas_format(caplog):
    s = pl.Series("x", [1, 2, 3, 4, 5])
    with caplog.at_level(logging.INFO, logger="raffalib.polars"):
        s.raffa.startlog()
        s = s.filter(s > 2)
        s.raffa.endlog(timeit=False)
    assert any("Removed 2/5 (40.00%) values." in r.message for r in caplog.records)


def test_endlog_timeit_appends_duration(caplog):
    df = pl.DataFrame({"a": [1, 2, 3]})
    with caplog.at_level(logging.INFO, logger="raffalib.polars"):
        df.raffa.startlog()
        df.raffa.endlog(timeit=True)
    assert any("Took:" in r.message for r in caplog.records)


def test_join_inner_drops_source_columns(caplog):
    left = pl.DataFrame({"k": [1, 2, 3], "lv": [10, 20, 30]})
    right = pl.DataFrame({"k": [2, 3, 4], "rv": [200, 300, 400]})
    with caplog.at_level(logging.INFO, logger="raffalib.polars"):
        out = left.raffa.join(right, on="k", how="inner")
    assert out.shape[0] == 2
    assert set(out.columns) == {"k", "lv", "rv"}


def test_join_keep_row_index_requires_permute():
    pytest.importorskip("polars_permute")
    left = pl.DataFrame({"k": [1, 2, 3], "lv": [10, 20, 30]})
    right = pl.DataFrame({"k": [2, 3, 4], "rv": [200, 300, 400]})
    out = left.raffa.join(right, on="k", how="inner", keep_row_index=True)
    assert "source_left" in out.columns
    assert "source_right" in out.columns


def test_endlog_add_rows_to_empty_frame_is_guarded(caplog):
    df = pl.DataFrame({"a": []}, schema={"a": pl.Float64})
    with caplog.at_level(logging.INFO, logger="raffalib.polars"):
        df.raffa.startlog()
        df = df.vstack(pl.DataFrame({"a": [1.0, 2.0, 3.0]}))
        df.raffa.endlog(timeit=False)
    assert any("Added 3/0 (N/A) rows." in r.message for r in caplog.records)


def test_join_empty_output_does_not_divide_by_zero(caplog):
    left = pl.DataFrame({"k": [1, 2]})
    right = pl.DataFrame({"k": [3, 4]})
    with caplog.at_level(logging.INFO, logger="raffalib.polars"):
        out = left.raffa.join(right, on="k", how="inner")
    assert out.shape[0] == 0
    assert any("From both: 0/0 (N/A)" in r.message for r in caplog.records)


def test_replace_string_with_null():
    df = pl.DataFrame({"a": ["x", "NA", "y"]})
    out = df.raffa.replace_string_with_null("NA")
    assert out["a"].to_list() == ["x", None, "y"]


def test_to_docx_applies_heading_and_table_options(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    df = pl.DataFrame({"a": [1, 2], "b": ["AAA", "BBB"]})
    out = tmp_path / "styled.docx"
    df.raffa.to_docx(out, heading_text="Table 1", table_style="Light Grid")
    reopened = Document(str(out))
    assert "Table 1" in [p.text for p in reopened.paragraphs]
    assert reopened.tables[0].style.name == "Light Grid"
