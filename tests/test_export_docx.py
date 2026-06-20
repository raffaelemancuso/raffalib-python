import pytest

pytest.importorskip("docx")

from docx import Document  # noqa: E402

from raffalib.export_docx import DocxFile  # noqa: E402


def test_create_table_and_save(tmp_path):
    doc = DocxFile()
    doc.add_table(3, 2)
    doc.table.cell(0, 0).text = "header"
    doc.table.cell(1, 1).text = "value"
    out = tmp_path / "out.docx"
    doc.save(out)

    assert out.exists()
    reopened = Document(str(out))
    assert len(reopened.tables) == 1
    table = reopened.tables[0]
    assert len(table.rows) == 3
    assert len(table.columns) == 2
    assert table.cell(0, 0).text == "header"
    assert table.cell(1, 1).text == "value"


def test_portrait_is_default():
    doc = DocxFile()
    assert doc.page_width < doc.page_height


def test_landscape_orientation():
    doc = DocxFile(landscape=True)
    assert doc.page_width > doc.page_height


def test_heading_written_to_document(tmp_path):
    doc = DocxFile(heading_text="My Heading")
    out = tmp_path / "heading.docx"
    doc.save(out)

    reopened = Document(str(out))
    texts = [p.text for p in reopened.paragraphs]
    assert "My Heading" in texts


def test_save_without_table(tmp_path):
    doc = DocxFile()
    out = tmp_path / "empty.docx"
    doc.save(out)
    assert out.exists()


def test_with_table_routes_init_and_table_kwargs(tmp_path):
    # heading_text -> __init__, table_style -> add_table; both must apply
    doc = DocxFile.with_table(2, 2, heading_text="My Heading", table_style="Light Grid")
    out = tmp_path / "routed.docx"
    doc.save(out)

    reopened = Document(str(out))
    assert "My Heading" in [p.text for p in reopened.paragraphs]
    assert len(reopened.tables) == 1
    assert reopened.tables[0].style.name == "Light Grid"


def test_with_table_rejects_unknown_kwarg():
    with pytest.raises(TypeError, match="unexpected keyword argument 'bogus'"):
        DocxFile.with_table(2, 2, bogus=123)
