# raffalib-python Miscellaneous functions
# Copyright (C) 2026 Raffaele Mancuso
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

"""Utilities for creating and exporting Word documents (.docx)."""

import inspect
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from io import BytesIO


class DocxFile:
    """
    A class for creating Word (.docx) documents with tables and figures.

    :ivar doc: The underlying python-docx Document object.
    :ivar page_height: Page height in millimeters (default A4: 210).
    :ivar page_width: Page width in millimeters (default A4: 297).
    :ivar table: The table object if one has been added.

    :param page_height: Page height in millimeters. Defaults to 210 (A4).
    :type page_height: float
    :param page_width: Page width in millimeters. Defaults to 297 (A4).
    :type page_width: float
    :param landscape: If True, use landscape orientation. Defaults to False.
    :type landscape: bool
    :param left_margin: Left margin in millimeters. Defaults to 25.4.
    :type left_margin: float
    :param right_margin: Right margin in millimeters. Defaults to 25.4.
    :type right_margin: float
    :param top_margin: Top margin in millimeters. Defaults to 25.4.
    :type top_margin: float
    :param bottom_margin: Bottom margin in millimeters. Defaults to 25.4.
    :type bottom_margin: float
    :param heading_text: Optional text for the document heading. Defaults to None.
    :type heading_text: str | None
    :param heading_font_name: Font name for heading. Defaults to "Aptos".
    :type heading_font_name: str
    :param heading_font_size: Font size for heading in points. Defaults to 12.
    :type heading_font_size: int
    :param heading_bold: Whether heading is bold. Defaults to True.
    :type heading_bold: bool
    :param heading_italic: Whether heading is italic. Defaults to False.
    :type heading_italic: bool
    :param heading_underline: Whether heading is underlined. Defaults to False.
    :type heading_underline: bool
    :param heading_color: RGB color tuple for heading. Defaults to black (0, 0, 0).
    :type heading_color: tuple[int, int, int]
    :param heading_space_before: Space before heading in points. Defaults to 0.
    :type heading_space_before: int
    :param heading_space_after: Space after heading in points. Defaults to 6.
    :type heading_space_after: int
    """

    def __init__(
        self,
        page_height: float = 210,
        page_width: float = 297,
        landscape: bool = False,
        left_margin: float = 25.4,
        right_margin: float = 25.4,
        top_margin: float = 25.4,
        bottom_margin: float = 25.4,
        heading_text: str | None = None,
        heading_font_name: str = "Aptos",
        heading_font_size: int = 12,
        heading_bold: bool = True,
        heading_italic: bool = False,
        heading_underline: bool = False,
        heading_color: tuple[int, int, int] = (0x00, 0x00, 0x00),
        heading_space_before: int = 0,
        heading_space_after: int = 6,
    ):

        # create new document
        self.doc = Document()

        # Page layout
        section = self.doc.sections[0]
        if landscape:
            self.page_height = page_height
            self.page_width = page_width
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            self.page_height = page_width
            self.page_width = page_height
            section.orientation = WD_ORIENT.PORTRAIT

        section.page_height = Mm(self.page_height)
        section.page_width = Mm(self.page_width)

        # Page margins
        section.left_margin = Mm(left_margin)
        section.right_margin = Mm(right_margin)
        section.top_margin = Mm(top_margin)
        section.bottom_margin = Mm(bottom_margin)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        # Add section heading
        if heading_text:
            # Add heading
            self.doc.add_heading(heading_text)

            # Set style of the header
            style = self.doc.styles["Heading 1"]

            font = style.font

            font.name = heading_font_name
            # Setting the style of the header requires this additional workaround
            # see: https://stackoverflow.com/a/60922725/1719931
            rFonts = style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), heading_font_name)

            font.size = Pt(heading_font_size)
            font.bold = heading_bold
            font.italic = heading_italic
            font.underline = heading_underline

            font.color.rgb = RGBColor(*heading_color)

            paragraph_format = style.paragraph_format
            paragraph_format.space_before = Pt(heading_space_before)
            paragraph_format.space_after = Pt(heading_space_after)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        self.table = None

    def add_table(
        self,
        n_rows,
        n_cols,
        table_autofit: bool = True,
        table_style: str = "Table Grid",
        table_font_name: str = "Aptos",
        table_font_size: int = 12,
        table_header_font_name: str = "Aptos",
        table_header_font_size: int = 12,
        table_header_font_bold: bool = True,
    ):
        """
        Add a table to the document.

        :param n_rows: Number of rows in the table.
        :type n_rows: int
        :param n_cols: Number of columns in the table.
        :type n_cols: int
        :param table_autofit: Whether to auto-fit column widths. Defaults to True.
        :type table_autofit: bool
        :param table_style: Style name for the table. Defaults to "Table Grid".
        :type table_style: str
        :param table_font_name: Font name for table cells. Defaults to "Aptos".
        :type table_font_name: str
        :param table_font_size: Font size for table cells. Defaults to 12.
        :type table_font_size: int
        :param table_header_font_name: Font name for header row. Defaults to "Aptos".
        :type table_header_font_name: str
        :param table_header_font_size: Font size for header row. Defaults to 12.
        :type table_header_font_size: int
        :param table_header_font_bold: Whether header row is bold. Defaults to True.
        :type table_header_font_bold: bool
        """

        # Create table
        self.table_autofit = table_autofit
        self.table = self.doc.add_table(n_rows, n_cols)

        # Set table style
        # See: https://github.com/python-openxml/python-docx/issues/9
        self.table.style = table_style

        # Create style for table header (first row)
        style = self.doc.styles.add_style("CellHeader", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self.doc.styles["Normal"]
        font = style.font
        font.name = table_header_font_name
        font.size = Pt(table_header_font_size)
        font.bold = table_header_font_bold

        # Create style for table cells (second row onwards)
        style = self.doc.styles.add_style("CellText", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self.doc.styles["Normal"]
        font = style.font
        font.name = table_font_name
        font.size = Pt(table_font_size)

    @classmethod
    def with_table(cls, n_rows: int, n_cols: int, **kwargs) -> "DocxFile":
        """
        Build a :class:`DocxFile` and add a table to it in a single call.

        Each keyword argument is routed to either :meth:`__init__` (document- and
        heading-level options such as ``landscape`` or ``heading_text``) or
        :meth:`add_table` (table-level options such as ``table_style`` or
        ``table_font_size``), based on which method accepts it.

        :param n_rows: Number of rows in the table.
        :type n_rows: int
        :param n_cols: Number of columns in the table.
        :type n_cols: int
        :param kwargs: Options forwarded to :meth:`__init__` or :meth:`add_table`.
        :return: A DocxFile with the table created and accessible via ``.table``.
        :rtype: DocxFile
        :raises TypeError: If a keyword argument is accepted by neither method.
        """
        init_params = set(inspect.signature(cls.__init__).parameters) - {"self"}
        table_params = set(inspect.signature(cls.add_table).parameters) - {
            "self",
            "n_rows",
            "n_cols",
        }
        doc_kwargs: dict = {}
        table_kwargs: dict = {}
        for key, value in kwargs.items():
            if key in init_params:
                doc_kwargs[key] = value
            elif key in table_params:
                table_kwargs[key] = value
            else:
                raise TypeError(
                    f"DocxFile.with_table() got an unexpected keyword argument {key!r}"
                )
        doc = cls(**doc_kwargs)
        doc.add_table(n_rows, n_cols, **table_kwargs)
        return doc

    # Additional workaround to make autofit actually work
    # See: https://github.com/python-openxml/python-docx/issues/209#issuecomment-566128709
    def _table_autofit_hotfix(self):
        """
        Workaround to make table autofit actually work.

        This is an internal method that applies a fix for the python-docx autofit issue.
        """
        for column in self.table.columns:
            for cell in column.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.get_or_add_tcW()
                tcW.type = "auto"

    def add_matplotlib(self, fig, width_perc: float = 0.7):
        """
        Add a matplotlib figure to the document.

        :param fig: A matplotlib Figure object.
        :param width_perc: Width of the image as a percentage of page width. Defaults to 0.7.
        :type width_perc: float
        :return: Self for method chaining.
        :rtype: DocxFile
        """
        memfile = BytesIO()
        fig.savefig(memfile, format="png")
        self.doc.add_picture(memfile, width=Mm(self.page_width * width_perc))
        memfile.close()
        return self

    def add_plotly(
        self,
        fig,
        docx_width: float | None = None,
        docx_height: float | None = None,
        img_width: float | None = None,
        img_height: float | None = None,
        img_scale: float | None = None,
    ):
        """
        Add a plotly figure to the document.

        :param fig: A plotly Figure object.
        :param docx_width: Width of the image in the document, in millimeters. Defaults to None.
        :type docx_width: float | None
        :param docx_height: Height of the image in the document, in millimeters. Defaults to None.
        :type docx_height: float | None
        :param img_width: Rendered image width in logical pixels. Defaults to None.
        :type img_width: float | None
        :param img_height: Rendered image height in logical pixels. Defaults to None.
        :type img_height: float | None
        :param img_scale: Scale factor for the rendered image's physical resolution
            (>1 increases, <1 decreases). Defaults to None.
        :type img_scale: float | None
        :return: Self for method chaining.
        :rtype: DocxFile
        """
        memfile = BytesIO()
        # In addition to the image format, the to_image and write_image functions provide
        # arguments to specify the image width and height in logical pixels.
        # They also provide a scale parameter that can be used to increase (scale > 1)
        # or decrease (scale < 1) the physical resolution of the resulting image.
        fig.write_image(
            memfile, format="png", width=img_width, height=img_height, scale=img_scale
        )
        if docx_width:
            docx_width = Mm(docx_width)
        if docx_height:
            docx_height = Mm(docx_height)
        self.doc.add_picture(memfile, width=docx_width, height=docx_height)
        memfile.close()
        return self

    def save(self, outfp: Path):
        """
        Save the document to a file.

        :param outfp: Path to the output .docx file.
        :type outfp: Path
        """

        if self.table:
            # Hotfix for table autofit
            if self.table_autofit:
                self._table_autofit_hotfix()

            # Set table text style for header (row 0)
            for cell in self.table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = "CellHeader"

            # Set table text style for data cells (rows >= 1)
            for row in self.table.rows[1:]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = "CellText"  # assuming style named "Cell Text"

        # save the doc
        self.doc.save(outfp)
